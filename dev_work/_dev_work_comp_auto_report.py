
# IMPORTS ###################################################################################################################################

from bs4 import BeautifulSoup
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from fake_useragent import UserAgent
from math import radians, cos, sin, asin, sqrt
from openai import OpenAI
from pyppeteer import launch, errors as pyppeteer_errors
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from tqdm import tqdm
from transformers import MarianMTModel, MarianTokenizer
from urllib.parse import urlparse, urljoin
import asyncio # new
import certifi
import difflib
import google.generativeai as genai
import json
import logging
import numpy as np
import os
import pandas as pd
import pyautogui
import pytz
import queue
import random
import re
import requests
import speech_recognition as sr
import ssl
import subprocess
import threading
import time
import tkinter as tk
import traceback
import webbrowser
import wikipedia
import wolframalpha
import yfinance as yf

# CUSTOM IMPORTS ##############################################################################################################################


# CONSTANTS ###################################################################################################################################

load_dotenv()
ACTIVATION_WORD = os.getenv('ACTIVATION_WORD', 'robot')
USER_PREFERRED_LANGUAGE = os.getenv('USER_PREFERRED_LANGUAGE', 'en')  # 2-letter lowercase
USER_PREFERRED_VOICE = os.getenv('USER_PREFERRED_VOICE', 'Evan')  # Daniel
USER_PREFERRED_NAME = os.getenv('USER_PREFERRED_NAME', 'User')  # Title case
USER_SELECTED_PASSWORD = os.getenv('USER_SELECTED_PASSWORD', 'None')  
USER_SELECTED_HOME_CITY = os.getenv('USER_SELECTED_HOME_CITY', 'None')  # Title case
USER_SELECTED_HOME_COUNTY = os.getenv('USER_SELECTED_HOME_COUNTY', 'None')  # Title case
USER_SELECTED_HOME_STATE = os.getenv('USER_SELECTED_HOME_STATE', 'None')  # Title case
USER_SELECTED_HOME_COUNTRY = os.getenv('USER_SELECTED_HOME_COUNTRY', 'None')  # 2-letter country code
USER_SELECTED_HOME_LAT = os.getenv('USER_SELECTED_HOME_LAT', 'None')  # Float with 6 decimal places
USER_SELECTED_HOME_LON = os.getenv('USER_SELECTED_HOME_LON', 'None')  # Float with 6 decimal places 
USER_SELECTED_TIMEZONE = os.getenv('USER_SELECTED_TIMEZONE', 'America/Chicago')  # Country/State format
USER_STOCK_WATCH_LIST = os.getenv('USER_STOCK_WATCH_LIST', 'None').split(',')  # Comma separated list of stock symbols
USER_DOWNLOADS_FOLDER = os.getenv('USER_DOWNLOADS_FOLDER')
PROJECT_VENV_DIRECTORY = os.getenv('PROJECT_VENV_DIRECTORY')
PROJECT_ROOT_DIRECTORY = os.getenv('PROJECT_ROOT_DIRECTORY')
SCRIPT_DIR_PATH = os.path.dirname(os.path.realpath(__file__))
SRC_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, 'src')
ARCHIVED_DEV_VERSIONS_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, '_archive')
FILE_DROP_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, 'app_generated_files')
LOCAL_LLMS_DIR = os.path.join(PROJECT_ROOT_DIRECTORY, 'app_local_models')
NOTES_DROP_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, 'app_base_knowledge')
SECRETS_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, '_secrets')
SOURCE_DATA_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, 'app_source_data')
SRC_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, 'src')
TESTS_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, '_tests')
UTILITIES_DIR_PATH = os.path.join(PROJECT_ROOT_DIRECTORY, 'utilities')
folders_to_create = [ARCHIVED_DEV_VERSIONS_PATH, FILE_DROP_DIR_PATH, LOCAL_LLMS_DIR, NOTES_DROP_DIR_PATH, SECRETS_DIR_PATH, SOURCE_DATA_DIR_PATH, SRC_DIR_PATH, TESTS_DIR_PATH, UTILITIES_DIR_PATH]
for folder in folders_to_create:
    if not os.path.exists(folder):
        os.makedirs(folder)

# def create_ssl_context():
#     return ssl.create_default_context(cafile=certifi.where())

# ssl._create_default_https_context = create_ssl_context
# context = create_ssl_context()
# print(f"""SSL Context Details: 
#     CA Certs File: {context.cert_store_stats()} 
#     Protocol: {context.protocol} 
#     Options: {context.options} 
#     Verify Mode: {context.verify_mode}
#     Verify Flags: {context.verify_flags}
#     Check Hostname: {context.check_hostname}
#     CA Certs Path: {certifi.where()}
#     """)

# Set API keys and other information from environment variables
open_weather_api_key = os.getenv('OPEN_WEATHER_API_KEY')
openai_api_key=os.getenv('OPENAI_API_KEY')
google_cloud_api_key = os.getenv('GOOGLE_CLOUD_API_KEY')
google_maps_api_key = os.getenv('GOOGLE_MAPS_API_KEY')
google_gemini_api_key = os.getenv('GOOGLE_GEMINI_API_KEY')

# Initialize LLM
genai.configure(api_key=google_gemini_api_key)
model = genai.GenerativeModel('gemini-pro')
vision_model = genai.GenerativeModel('gemini-pro-vision')

# Set the terminal output display options
pd.set_option('display.max_columns', 10)
pd.set_option('display.max_rows', 150)
pd.set_option('display.width', None)
pd.set_option('display.max_colwidth', 35)
pd.set_option('display.expand_frame_repr', False)
pd.set_option('display.precision', 1) 

# FUNCTION DEFINITIONS ###################################################################################################################################

address = "1486 East Valley Road Montecito, CA 93108"
search_distance = 10000
    
class AddressResearcher:
    def __init__(self, api_key):
        self.api_key = api_key
        self.geocode_url = "https://maps.googleapis.com/maps/api/geocode/json"
        self.place_search_url = "https://maps.googleapis.com/maps/api/place/nearbysearch/json"
        self.data = {}
        self.type_groupings = {
            'restaurants': ['restaurant'],
        }

    USER_AGENTS = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.0.3 Safari/605.1.15',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:68.0) Gecko/20100101 Firefox/68.0',
        'Mozilla/5.0 (Windows NT 10.0; WOW64; Trident/7.0; rv:11.0) like Gecko',
        'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:88.0) Gecko/20100101 Firefox/88.0',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_14_6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
        'Mozilla/5.0 (iPhone; CPU iPhone OS 14_4 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.0 Mobile/15E148 Safari/604.1',
        'Mozilla/5.0 (Linux; Android 10; SM-G960U) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.181 Mobile Safari/537.36',
        'Mozilla/5.0 (iPad; CPU OS 14_4 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.0 Mobile/15E148 Safari/604.1',
        'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.96 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.212 Safari/537.36',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/13.1 Safari/605.1.15',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:78.0) Gecko/20100101 Firefox/78.0',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.75 Safari/537.36',
        'Mozilla/5.0 (iPhone; CPU iPhone OS 13_3 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148',
        'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/85.0.4183.121 Safari/537.36',
        'Mozilla/5.0 (X11; Ubuntu; Linux i686; rv:64.0) Gecko/20100101 Firefox/64.0',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/604.1 (KHTML, like Gecko) Edge/18.19582',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_3) AppleWebKit/537.75.14 (KHTML, like Gecko) Version/7.0.3 Safari/7046A194A',
        'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/77.0.3865.120 Safari/537.36',
        'Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.81 Safari/537.36',
        'Mozilla/5.0 (iPad; CPU OS 13_2 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148',
        'Mozilla/5.0 (Macintosh; U; Intel Mac OS X 10_6_6; en-us) AppleWebKit/533.20.25 (KHTML, like Gecko) Version/5.0.4 Safari/533.20.27',
        'Mozilla/5.0 (Windows NT 10.0) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.102 Safari/537.36 Edge/18.19577',
        'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/601.1 (KHTML, like Gecko) Chrome/50.0.2661.102 Safari/601.1',
        'Mozilla/5.0 (Windows NT 6.2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36',
        'Mozilla/5.0 (iPhone; CPU iPhone OS 12_4 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148',
        'Mozilla/5.0 (Windows NT 6.1; Win64; x64; rv:57.0) Gecko/20100101 Firefox/57.0',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_6) AppleWebKit/601.7.8 (KHTML, like Gecko) Version/9.1.2 Safari/601.7.7',
        'Mozilla/5.0 (X11; Fedora; Linux x86_64; rv:61.0) Gecko/20100101 Firefox/61.0',
    ]
    
    ### INITIAL PASS OF REPORT FROM API CALL ###

    def perform_research(self, address):
        distance = search_distance
        print(f"Starting research for address: {address}")
        lat, lng = self.geocode_address(address)
        if lat is None or lng is None:
            print("Geocoding failed.")
            return

        all_places = self.search_nearby(lat, lng, distance)

        # Grouping and sorting
        grouped_places = {}
        for place in all_places:
            group = place['group']
            place['weighted_score'] = self.calculate_weighted_score(place)
            if group not in grouped_places:
                grouped_places[group] = []
            grouped_places[group].append(place)

        for group in grouped_places:
            grouped_places[group].sort(key=lambda x: x['weighted_score'], reverse=True)

        self.data = {'address': address, 'grouped_places': grouped_places}
    
    def geocode_address(self, address):
        params = {'address': address, 'key': self.api_key}
        response = requests.get(self.geocode_url, params=params)
        if response.status_code != 200:
            print(f"Geocode API request failed with status code: {response.status_code}")
            return None, None
        geocode_data = response.json()
        if geocode_data['status'] != 'OK':
            print(f"Geocode API response error: {geocode_data['status']}")
            return None, None
        location = geocode_data['results'][0]['geometry']['location']
        return location['lat'], location['lng']

    def haversine(self, lon1, lat1, lon2, lat2):
        """
        Calculate the great circle distance in kilometers between two points 
        on the earth (specified in decimal degrees)
        """
        # convert decimal degrees to radians 
        lon1, lat1, lon2, lat2 = map(radians, [lon1, lat1, lon2, lat2])

        # Haversine formula 
        dlon = lon2 - lon1 
        dlat = lat2 - lat1 
        a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
        c = 2 * asin(sqrt(a)) 
        km = 6371 * c # Radius of earth in kilometers
        return km

    def search_nearby(self, lat, lng, radius):
        all_places = []
        missing_coordinate_places = []

        for group, types in self.type_groupings.items():
            for place_type in types:
                next_page_token = None
                while True:
                    params = {
                        'location': f"{lat},{lng}",
                        'radius': radius,
                        'type': place_type,
                        'key': self.api_key,
                        'pagetoken': next_page_token
                    }
                    response = requests.get(self.place_search_url, params=params)
                    if response.status_code == 200:
                        search_data = response.json()
                        if search_data['status'] == 'OK':
                            for place in search_data['results']:
                                place_details = self.get_place_details(place['place_id'])
                                if place_details:
                                    place_lat = place_details.get('geometry', {}).get('location', {}).get('lat')
                                    place_lng = place_details.get('geometry', {}).get('location', {}).get('lng')
                                    if place_lat and place_lng:
                                        distance = self.haversine(lng, lat, place_lng, place_lat)
                                        place_details['distance_from_target_address'] = distance
                                    else:
                                        missing_coordinate_places.append(place_details)
                                    place_details['group'] = group
                                    all_places.append(place_details)
                            next_page_token = search_data.get('next_page_token')
                            if not next_page_token:
                                break
                            time.sleep(2)  # Delay for token validity
                        else:
                            break
                    else:
                        break

        # Process the batch of missing coordinate places
        self.process_missing_coordinate_places(missing_coordinate_places, lat, lng)

        return all_places

    def process_missing_coordinate_places(self, missing_places, origin_lat, origin_lng):
        # Google Maps Distance Matrix API endpoint
        distance_matrix_url = "https://maps.googleapis.com/maps/api/distancematrix/json"

        for place in missing_places:
            # Building the address from the place details
            destination_address = place.get('vicinity', 'Unknown Location')
            params = {
                'origins': f"{origin_lat},{origin_lng}",
                'destinations': destination_address,
                'key': self.api_key
            }
            response = requests.get(distance_matrix_url, params=params)
            if response.status_code == 200:
                matrix_data = response.json()
                if matrix_data['status'] == 'OK':
                    # Extracting distance information
                    result = matrix_data['rows'][0]['elements'][0]
                    if result['status'] == 'OK':
                        distance_value = result['distance']['value'] / 1000  # Convert meters to kilometers
                        place['distance_from_target_address'] = distance_value
                    else:
                        print(f"Distance calculation failed for {destination_address}: {result['status']}")
                else:
                    print(f"Distance Matrix API error: {matrix_data['status']}")
            else:
                print(f"Distance Matrix API request failed with status code: {response.status_code}")
    
    def calculate_weighted_score(self, place):
        rating = place.get('rating', 0)
        price_level = place.get('price_level', 0)
        distance = place.get('distance_from_target_address', float('inf'))
        distance_factor = 1 / (distance + 1)  # Add 1 to avoid division by zero
        return (rating * 0.4 + price_level * 0.6) * distance_factor

    def get_place_details(self, place_id):
        print(f"Fetching details for place ID: {place_id}")
        details_url = "https://maps.googleapis.com/maps/api/place/details/json"
        params = {
            'place_id': place_id,
            'fields': 'name,place_id,url,type,opening_hours,utc_offset,reservable,rating,price_level,review,user_ratings_total,formatted_phone_number,international_phone_number,website,scope,vicinity,formatted_address,business_status,dine_in,editorial_summary,serves_breakfast,serves_brunch,serves_dinner,serves_lunch,serves_wine,serves_vegetarian_food,takeout',
            'key': self.api_key
        }
        response = requests.get(details_url, params=params)
        if response.status_code == 200:
            details_data = response.json()
            if details_data['status'] == 'OK':
                return details_data['result']
        return {}

    def save_research_report(self, filename):
        with open(filename, 'w') as file:
            json.dump(self.data, file, indent=4)

    ### WEB SCRAPING METHODS TO APPEND MORE DATA TO THE REPORT DATASET ###

    def scrape_website_data(self, url):
        print(f"Scraping website: {url}")
        try:
            response = requests.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')

            # Base URL extraction for concatenating with relative links
            base_url = '{uri.scheme}://{uri.netloc}'.format(uri=urlparse(url))

            # Extract and filter links
            links = soup.find_all('a', href=True)
            # menu_links = set(self.get_full_link(link['href'], base_url) for link in links if 'menu' in link.get_text().lower() and self.is_valid_link(link['href']))
            menu_links = set(self.get_full_link(link['href'], base_url) for link in links if any(keyword in link.get_text().lower() for keyword in ['menu', 'dinner', 'lunch', 'food', 'wine', 'drink', 'drinks', 'eat', 'cocktail', 'cocktails', 'bar', 'bites', 'snack', 'snacks',]) and self.is_valid_link(link['href']))
            event_links = set(self.get_full_link(link['href'], base_url) for link in links if any(word in link.get_text().lower() for word in ['event', 'calendar', 'schedule',]) and self.is_valid_link(link['href']))
            pdf_links = set(self.get_full_link(link['href'], base_url) for link in links if '.pdf' in link['href'] and self.is_valid_link(link['href']))

            return {'menu_links': list(menu_links), 'event_links': list(event_links), 'pdf_links': list(pdf_links)}

        except requests.exceptions.Timeout:
            print(f"Timeout occurred while scraping {url}")
            return {}
        except requests.exceptions.RequestException as e:
            print(f"Error scraping {url}: {e}")
            return {}

    def get_full_link(self, link, base_url):
        if link.startswith('/'):
            return urljoin(base_url, link)
        return link

    def is_valid_link(self, link):
        # Add logic to filter out invalid or duplicate links
        return link.startswith("http") or (link.startswith("/") and not link.startswith("//"))
        
    def augment_place_details_with_web_data(self):
        print("Augmenting place details with web data")
        for group, places in self.data['grouped_places'].items():
            for place in tqdm(places, desc=f"Processing {group}"):
                if 'website' in place:
                    web_data = self.scrape_website_data(place['website'])
                    place.update(web_data)
                    time.sleep(2)  
                    
        time.sleep(2) 
        
    def generate_summary(self):
        for group, places in self.data['grouped_places'].items():
            for place in tqdm(places, desc=f"Generating in-market comp summaries for {group}"):
                try:
                    place_data_str = json.dumps(place, ensure_ascii=False)
                    prompt = (
                        f"""### <SYSTEM MESSAGE> <START> ### 
                        Gemini, you are a hospitality, restaurant, and F&B market comparison business analyst. 
                        The audience for your output is is an operations research analyst who is attempting to draw conclusions 
                        about the affect that various businesses may have on their potential future business in the same market. 
                        The new business will be an American fine dining restaurant. 
                        The data shown below is for a business that is located near the aformentioned luxury American fine dining restaurant. 
                        You will write a *CONCISE AND SHORT* summary, following the instructions below. 
                        Your summary must be only a few simple sentences. 
                        List the details about this business that would be relevant to a hospitality and restaurant and F&B market analyst. 
                        The goal is to examine this business within the context of its local market (geographic and economic). 
                        Explain everything you know from your training data about this business at this specific location. 
                        List important headlines and info about this location (if any). 
                        This business is most likely a restaurant. 
                        List any headlines, awards or accolades this business has recieved such as Michelin stars, James Beard awards, etc. 
                        Provide any known information about the business such as whether they have any Michelin stars, have been in the news, if they have a famous chef, etc. 
                        The business in this dataset provided is located near an American fine dining restaurant. 
                        You will offer your *market business insights* about the business listed below as it would relate to a nearby American fine dining restaurant. 
                        The nearby comparison restaurant is a luxury American fine dining establishment with a robust wine list and elevated style of service. 
                        Your goal is to explain the potential interplay and nuances of the relationship between these two entities. 
                        Write a short and concise summary of the business in the provided data, and how this business will affect the American fine dining restaurant - either positively, negatively, or neither. 
                        We want to learn if the restaurant shown in this data will help or hinder the nearby American fine dining restaurant's business. 
                        The comparison restaurant is an American fine dining restaurant, located near the business in the provided data. 
                        The distance between the businesses is listed in the data. 
                        The goal is to understand the relationships and potential competition level in the area of the American fine dining restaurant and draw business hypotheses about the potential relationships between the restaurant and the other entities in the area, such as the one shown in this data. 
                        Write a summary of how this business could affect the nearby American fine dining restaurant. 
                        Explain everything about this business as a potential competitor or significant entity within the American restaurant's local market. 
                        If the business is a restaurant, describe their menu offerings in detail and whether or not they would directly compete with a nearby American fine dining restaurant. 
                        Only analyze how the business in this data may passively influence a nearby American fine dining restaurant. Do not give any recommendations or advice about the two entities partnering directly. 
                        Also include any of your own knowledge about the listed business and its relevance to a restaurant located nearby within the same market. 
                        Your output must be a short and concise summary of the business in the context of its local market, with your own thoughtful insights included in the summary. 
                        Keep it simple and concise and only include factual data. 
                        Please limit your responses to a few sentences, but include substantial information. 
                        Make sure your response is simple and easy to read quickly. 
                        Limit your response to no more than 1 paragraph of direct and concise text. 
                        Be consice and direct. 
                        Here is the object to analyze: {place_data_str}
                        ### <SYSTEM MESSAGE> <END> ###"""
                    )
                    response = model.generate_content(prompt)
                    summary = response.text
                    place['gemini_summary'] = summary
                    time.sleep(2) 
                except Exception as e:
                    print(f"Error generating summary for {place.get('name', 'N/A')}: {e}")
                    place['gemini_summary'] = 'N/A'
                    time.sleep(2)

        # Optionally return the whole data if needed
        return self.data
    
    ### FILE SAVING ###

    # def save_report_as_word(self, filename):
    #     doc = Document()
    #     style = doc.styles['Normal']
    #     font = style.font
    #     font.name = 'Open Sans'
    #     font.size = Pt(12)

    #     doc.add_heading('Research Report', level=1)
    #     doc.add_paragraph(f'Address: {self.data["address"]}\n')
        
    #     self.add_summary_to_doc(doc)

    #     for group, places in self.data['grouped_places'].items():
    #         doc.add_heading(f'{group.capitalize()}:', level=2)
    #         for place in places:
    #             if isinstance(place, dict):
    #                 self.add_place_details_to_doc(doc, place)
    #             else:
    #                 print(f"Warning: Expected a dictionary for 'place', got {type(place)}")

    #     doc.save(filename)

    # def add_place_details_to_doc(self, doc, place):
    #     # Ensure 'place' is a dictionary
    #     if not isinstance(place, dict):
    #         print(f"Warning: Expected a dictionary for 'place', got {type(place)}")
    #         return

    #     doc.add_heading(place.get('name', 'N/A'), level=3)
                
    #     doc.add_paragraph(f"Distance from Address: {place.get('distance_from_target_address', 'N/A')} km", style='List Bullet')
        
    #     editorial_summary = place.get('editorial_summary', {})
    #     summary_str = editorial_summary.get('overview', 'N/A')
    #     doc.add_paragraph(f"Editorial summary: {summary_str}", style='List Bullet')

    #     types = ', '.join(place.get('types', ['N/A']))
    #     doc.add_paragraph(f"Types: {types}", style='List Bullet')
        
    #     hours = place.get('opening_hours', {})
    #     if 'weekday_text' in hours:
    #         hours_str = '\n'.join(hours['weekday_text'])
    #     else:
    #         hours_str = 'N/A'
    #     doc.add_paragraph(f"Hours of operation: {hours_str}", style='List Bullet')

    #     table = doc.add_table(rows=1, cols=2)
    #     hdr_cells = table.rows[0].cells
    #     hdr_cells[0].text = 'Category'
    #     hdr_cells[1].text = 'Data'
    
    #     # Standard fields
    #     standard_fields = ['price_level', 'rating', 'user_ratings_total', 'business_status', 'website', 'url', 'international_phone_number', 'formatted_address']
    #     for field in standard_fields:
    #         row_cells = table.add_row().cells
    #         row_cells[0].text = field.replace('_', ' ').capitalize()
    #         row_cells[1].text = str(place.get(field, 'N/A'))

    #     # Additional boolean fields to convert to Yes/No for the report
    #     additional_fields = ['dine_in', 'takeout', 'reservable', 'serves_breakfast', 'serves_brunch', 'serves_lunch', 'serves_dinner', 'serves_wine', 'serves_vegetarian_food']
    #     for field in additional_fields:
    #         value = place.get(field, 'N/A')
    #         value_text = 'Yes' if value is True else 'No' if value is False else str(value)
    #         row_cells = table.add_row().cells
    #         row_cells[0].text = field.replace('_', ' ').capitalize()
    #         row_cells[1].text = value_text

    #     # List fields
    #     self.add_list_to_doc(doc, 'Menu Links', place.get('menu_links', []))
    #     self.add_list_to_doc(doc, 'Event Links', place.get('event_links', []))
    #     self.add_list_to_doc(doc, 'PDF Links', place.get('pdf_links', []))
        
    #     doc.add_paragraph(f"AI Summary: {place.get('gemini_summary', 'N/A')}", style='List Bullet')
            
    # def add_list_to_doc(self, doc, title, items):
    #     if items:
    #         doc.add_paragraph(f'{title}:', style='List Bullet')
    #         for item in items:
    #             # Create a sub-bullet point for each item
    #             p = doc.add_paragraph(style='List Bullet 2')
    #             p.add_run(item)

    # def add_summary_to_doc(self, doc):
    #     doc.add_heading('Summary of Findings', level=2)
    #     for group, places in self.data['grouped_places'].items():
    #         doc.add_heading(f'{group.capitalize()}:', level=3)
            
    #         # Create a table for this group with an additional column for the weighted score
    #         table = doc.add_table(rows=1, cols=5)
    #         hdr_cells = table.rows[0].cells
    #         hdr_cells[0].text = 'Weighted Match Score'
    #         hdr_cells[1].text = 'Name'
    #         hdr_cells[2].text = 'Rating'
    #         hdr_cells[3].text = 'Price Level'
    #         hdr_cells[4].text = 'Distance (km)'

    #         for place in places:
    #             row_cells = table.add_row().cells
    #             row_cells[0].text = str(round(place.get('weighted_score', 0), 2))  # Weighted score rounded to 2 decimal places
    #             row_cells[1].text = place.get('name', 'N/A')
    #             row_cells[2].text = str(place.get('rating', 'N/A'))
    #             row_cells[3].text = "$" * place.get('price_level', 0)
    #             row_cells[4].text = f"{round(place.get('distance_from_target_address', 0), 1)}"

    @staticmethod
    def format_weekday_text(hours):
        return '\n'.join(hours.get('weekday_text', ['N/A'])) if 'weekday_text' in hours else 'N/A'

    @staticmethod
    def format_reviews(reviews):
        formatted_reviews = []
        for review in reviews:
            # Convert Unix timestamp to human-readable date
            review_text = (f"Author: {review.get('author_name', 'Anonymous')}, "
                           f"Rating: {review.get('rating', 'N/A')}, "
                           f"Review: {review.get('text', 'No review text provided')}")
            formatted_reviews.append(review_text)
        return '\n'.join(formatted_reviews) if formatted_reviews else 'N/A'
    
    @staticmethod
    def format_links(links):
        return ', '.join(links) if links else 'N/A'

    @staticmethod
    def save_report_as_csv(json_file_path, csv_file_path):
        with open(json_file_path, 'r') as file:
            data = json.load(file)

        data_for_df = []
        for group, places in data['grouped_places'].items():
            for place in places:
                place_data = {
                    'Name': place.get('name', 'N/A'),
                    'Business Type': group.capitalize(),
                    'Distance (km)': round(place.get('distance_from_target_address', 0), 1),
                    'Weighted Relevance Score': round(place.get('weighted_score', 0), 2),
                    'Price Level': place.get('price_level', 'N/A'),
                    'Average Rating': place.get('rating', 'N/A'),
                    'Operating Hours': AddressResearcher.format_weekday_text(place.get('opening_hours', {})),
                    'Tags': ', '.join(place.get('types', ['N/A'])),
                    'Google Summary': place.get('editorial_summary', {}).get('overview', 'N/A'),
                    'AI Summary 1': place.get('gemini_summary', 'N/A'),
                    'Total Reviews': place.get('user_ratings_total', 'N/A'),
                    '5 Most Relevant Reviews': AddressResearcher.format_reviews(place.get('reviews', [])),
                    'Business Status': place.get('business_status', 'N/A'),
                    'Website': place.get('website', 'N/A'),
                    'Google Maps URL': place.get('url', 'N/A'),
                    'Menu Links': AddressResearcher.format_links(place.get('menu_links', [])),
                    'Event Links': AddressResearcher.format_links(place.get('event_links', [])),
                    'PDF Links': AddressResearcher.format_links(place.get('pdf_links', [])),
                    'Address': place.get('formatted_address', 'N/A'),
                    'Phone Number': place.get('formatted_phone_number', 'N/A'),
                    'International Phone Number': place.get('international_phone_number', 'N/A'),
                    'Reservable': 'Yes' if place.get('reservable') else 'No',
                    'Dine In': 'Yes' if place.get('dine_in') else 'No',
                    'Takeout': 'Yes' if place.get('takeout') else 'No',
                    'Serves Breakfast': 'Yes' if place.get('serves_breakfast') else 'No',
                    'Serves Lunch': 'Yes' if place.get('serves_lunch') else 'No',
                    'Serves Dinner': 'Yes' if place.get('serves_dinner') else 'No',
                    'Serves Wine': 'Yes' if place.get('serves_wine') else 'No',
                    'Service Vegetarian Food': 'Yes' if place.get('serves_vegetarian_food') else 'No',
                    'Menu Content': place.get('menu_content', 'N/A')
                }
                data_for_df.append(place_data)

        df = pd.DataFrame(data_for_df)
        df.to_csv(csv_file_path, index=False)

    ### PRINTING THE FINDINGS TO THE CONSOLE ###
    
    @staticmethod
    def print_human_readable(json_file_path):
        def print_dict(d, indent=0):
            for key, value in d.items():
                if isinstance(value, dict):
                    print(' ' * indent + f"{key.capitalize()}:")
                    print_dict(value, indent + 4)
                elif isinstance(value, list):
                    if all(isinstance(item, dict) for item in value):
                        for item in value:
                            print_dict(item, indent + 4)
                    else:
                        print(' ' * indent + f"{key.capitalize()}: {', '.join(map(str, value))}")
                else:
                    print(' ' * indent + f"{key.replace('_', ' ').capitalize()}: {value}")

        with open(json_file_path, 'r') as file:
            data = json.load(file)

        print(f"Research Report for Address: {data['address']}\n")
        print_dict(data['grouped_places'])
        
if __name__ == '__main__':
    now = datetime.now()
    address_researcher = AddressResearcher(google_maps_api_key)
    address_researcher.perform_research(address)
    address_researcher.augment_place_details_with_web_data()
    address_researcher.generate_summary()
    AddressResearcher.print_human_readable(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented_{address}.json")
    
    address_researcher.save_research_report(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented_{address}.json")
    if os.path.exists(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented_{address}.json"):
        print("Research json saved successfully.")
    else:
        print("Research json save failed.")
        
    address_researcher.save_report_as_word(f"{FILE_DROP_DIR_PATH}/research_report_doc_formatted_{address}.docx")
    if os.path.exists(f"{FILE_DROP_DIR_PATH}/research_report_doc_formatted_{address}.docx"):
        print("Research doc saved successfully.")
    else:
        print("Research doc save failed.")
        
    AddressResearcher.save_report_as_csv(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented_{address}.json", f"{FILE_DROP_DIR_PATH}/research_report_spreadsheet_{address}.csv")
    if os.path.exists(f"{FILE_DROP_DIR_PATH}/research_report_spreadsheet_{address}.csv"):
        print("Research csv saved successfully.")
    else:    
        print("Research csv save failed.")


