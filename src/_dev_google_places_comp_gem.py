
# IMPORTS ###################################################################################################################################

from bs4 import BeautifulSoup
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from math import radians, cos, sin, asin, sqrt
from openai import OpenAI
from tqdm import tqdm
from transformers import MarianMTModel, MarianTokenizer
from urllib.parse import urlparse, urljoin
import certifi
import datetime
import google.generativeai as genai
import json
import numpy as np
import os
import pandas as pd
import pyautogui
import pytz
import queue
import re
import requests
import speech_recognition as sr
import ssl
import subprocess
import threading
import time
import tkinter as tk
import traceback
import webbrowser
import wikipedia
import wolframalpha
import yfinance as yf

# CUSTOM IMPORTS ##############################################################################################################################


# CONSTANTS ###################################################################################################################################

# bring in the environment variables from the .env file
load_dotenv()
ACTIVATION_WORD = os.getenv('ACTIVATION_WORD', 'robot')
USER_DOWNLOADS_FOLDER = os.getenv('USER_DOWNLOADS_FOLDER')
USER_PREFERRED_LANGUAGE = os.getenv('USER_PREFERRED_LANGUAGE', 'en')  # 2-letter lowercase
USER_PREFERRED_VOICE = os.getenv('USER_PREFERRED_VOICE', 'Daniel')  # Daniel
USER_PREFERRED_NAME = os.getenv('USER_PREFERRED_NAME', 'User')  # Title case
USER_SELECTED_PASSWORD = os.getenv('USER_SELECTED_PASSWORD', 'None')  
USER_SELECTED_HOME_CITY = os.getenv('USER_SELECTED_HOME_CITY', 'None')  # Title case
USER_SELECTED_HOME_COUNTY = os.getenv('USER_SELECTED_HOME_COUNTY', 'None')  # Title case
USER_SELECTED_HOME_STATE = os.getenv('USER_SELECTED_HOME_STATE', 'None')  # Title case
USER_SELECTED_HOME_COUNTRY = os.getenv('USER_SELECTED_HOME_COUNTRY', 'None')  # 2-letter country code
USER_SELECTED_HOME_LAT = os.getenv('USER_SELECTED_HOME_LAT', 'None')  # Float with 6 decimal places
USER_SELECTED_HOME_LON = os.getenv('USER_SELECTED_HOME_LON', 'None')  # Float with 6 decimal places 
USER_SELECTED_TIMEZONE = os.getenv('USER_SELECTED_TIMEZONE', 'America/Chicago')  # Country/State format
USER_STOCK_WATCH_LIST = os.getenv('USER_STOCK_WATCH_LIST', 'None').split(',')  # Comma separated list of stock symbols
PROJECT_VENV_DIRECTORY = os.getenv('PROJECT_VENV_DIRECTORY')

# establish relative file paths for the current script
SCRIPT_DIR_PATH = os.path.dirname(os.path.realpath(__file__))
PROJECT_ROOT_DIR_PATH = os.path.dirname(SCRIPT_DIR_PATH)
ARCHIVED_DEV_VERSIONS_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, '_archive')
FILE_DROP_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_generated_files')
LOCAL_LLMS_DIR = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_local_models')
NOTES_DROP_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_base_knowledge')
SOURCE_DATA_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_source_data')
TESTS_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, '_tests')
UTILITIES_DIR_PATH = os.path.join(SCRIPT_DIR_PATH, 'utilities')

folders_to_create = [ARCHIVED_DEV_VERSIONS_PATH, FILE_DROP_DIR_PATH, LOCAL_LLMS_DIR, NOTES_DROP_DIR_PATH, SOURCE_DATA_DIR_PATH, TESTS_DIR_PATH, UTILITIES_DIR_PATH]
for folder in folders_to_create:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Set API keys and other information from environment variables
open_weather_api_key = os.getenv('OPEN_WEATHER_API_KEY')
wolfram_app_id = os.getenv('WOLFRAM_APP_ID')
openai_api_key=os.getenv('OPENAI_API_KEY')
google_cloud_api_key = os.getenv('GOOGLE_CLOUD_API_KEY')
google_maps_api_key = os.getenv('GOOGLE_MAPS_API_KEY')
google_gemini_api_key = os.getenv('GOOGLE_GEMINI_API_KEY')

# Initialize LLM
genai.configure(api_key=google_gemini_api_key)
model = genai.GenerativeModel('gemini-pro')

# Set the terminal output display options
pd.set_option('display.max_columns', 10)
pd.set_option('display.max_rows', 150)
pd.set_option('display.width', None)
pd.set_option('display.max_colwidth', 35)
pd.set_option('display.expand_frame_repr', False)
pd.set_option('display.precision', 1) 

# FUNCTION DEFINITIONS ###################################################################################################################################

address = "1750 Redwood Hwy, Corte Madera, CA 94925"
search_distance = 10000
    
class AddressResearcher:
    def __init__(self, api_key):
        self.api_key = api_key
        self.geocode_url = "https://maps.googleapis.com/maps/api/geocode/json"
        self.place_search_url = "https://maps.googleapis.com/maps/api/place/nearbysearch/json"
        self.data = {}
        self.type_groupings = {
            'restaurants': ['restaurant'],
        }

    ### INITIAL PASS OF REPORT FROM API CALL ###

    def perform_research(self, address):
        distance = search_distance
        print(f"Starting research for address: {address}")
        lat, lng = self.geocode_address(address)
        if lat is None or lng is None:
            print("Geocoding failed.")
            return

        all_places = self.search_nearby(lat, lng, distance)

        # Grouping and sorting
        grouped_places = {}
        for place in all_places:
            group = place['group']
            place['weighted_score'] = self.calculate_weighted_score(place)
            if group not in grouped_places:
                grouped_places[group] = []
            grouped_places[group].append(place)

        for group in grouped_places:
            grouped_places[group].sort(key=lambda x: x['weighted_score'], reverse=True)

        self.data = {'address': address, 'grouped_places': grouped_places}
    
    def geocode_address(self, address):
        params = {'address': address, 'key': self.api_key}
        response = requests.get(self.geocode_url, params=params)
        if response.status_code != 200:
            print(f"Geocode API request failed with status code: {response.status_code}")
            return None, None
        geocode_data = response.json()
        if geocode_data['status'] != 'OK':
            print(f"Geocode API response error: {geocode_data['status']}")
            return None, None
        location = geocode_data['results'][0]['geometry']['location']
        return location['lat'], location['lng']

    def haversine(self, lon1, lat1, lon2, lat2):
        """
        Calculate the great circle distance in kilometers between two points 
        on the earth (specified in decimal degrees)
        """
        # convert decimal degrees to radians 
        lon1, lat1, lon2, lat2 = map(radians, [lon1, lat1, lon2, lat2])

        # Haversine formula 
        dlon = lon2 - lon1 
        dlat = lat2 - lat1 
        a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
        c = 2 * asin(sqrt(a)) 
        km = 6371 * c # Radius of earth in kilometers
        return km

    def search_nearby(self, lat, lng, radius):
        all_places = []
        missing_coordinate_places = []

        for group, types in self.type_groupings.items():
            for place_type in types:
                next_page_token = None
                while True:
                    params = {
                        'location': f"{lat},{lng}",
                        'radius': radius,
                        'type': place_type,
                        'key': self.api_key,
                        'pagetoken': next_page_token
                    }
                    response = requests.get(self.place_search_url, params=params)
                    if response.status_code == 200:
                        search_data = response.json()
                        if search_data['status'] == 'OK':
                            for place in search_data['results']:
                                place_details = self.get_place_details(place['place_id'])
                                if place_details:
                                    place_lat = place_details.get('geometry', {}).get('location', {}).get('lat')
                                    place_lng = place_details.get('geometry', {}).get('location', {}).get('lng')
                                    if place_lat and place_lng:
                                        distance = self.haversine(lng, lat, place_lng, place_lat)
                                        place_details['distance_from_target_address'] = distance
                                    else:
                                        missing_coordinate_places.append(place_details)
                                    place_details['group'] = group
                                    all_places.append(place_details)
                            next_page_token = search_data.get('next_page_token')
                            if not next_page_token:
                                break
                            time.sleep(2)  # Delay for token validity
                        else:
                            break
                    else:
                        break

        # Process the batch of missing coordinate places
        self.process_missing_coordinate_places(missing_coordinate_places, lat, lng)

        return all_places

    def process_missing_coordinate_places(self, missing_places, origin_lat, origin_lng):
        # Google Maps Distance Matrix API endpoint
        distance_matrix_url = "https://maps.googleapis.com/maps/api/distancematrix/json"

        for place in missing_places:
            # Building the address from the place details
            destination_address = place.get('vicinity', 'Unknown Location')
            params = {
                'origins': f"{origin_lat},{origin_lng}",
                'destinations': destination_address,
                'key': self.api_key
            }
            response = requests.get(distance_matrix_url, params=params)
            if response.status_code == 200:
                matrix_data = response.json()
                if matrix_data['status'] == 'OK':
                    # Extracting distance information
                    result = matrix_data['rows'][0]['elements'][0]
                    if result['status'] == 'OK':
                        distance_value = result['distance']['value'] / 1000  # Convert meters to kilometers
                        place['distance_from_target_address'] = distance_value
                    else:
                        print(f"Distance calculation failed for {destination_address}: {result['status']}")
                else:
                    print(f"Distance Matrix API error: {matrix_data['status']}")
            else:
                print(f"Distance Matrix API request failed with status code: {response.status_code}")
    
    def calculate_weighted_score(self, place):
        rating = place.get('rating', 0)
        price_level = place.get('price_level', 0)
        distance = place.get('distance_from_target_address', float('inf'))

        # Adjust weights as needed
        # Higher weight for price level and proximity, excluding user ratings total
        distance_factor = 1 / (distance + 1)  # Add 1 to avoid division by zero
        return (rating * 0.4 + price_level * 0.6) * distance_factor

    def get_place_details(self, place_id):
        print(f"Fetching details for place ID: {place_id}")
        details_url = "https://maps.googleapis.com/maps/api/place/details/json"
        params = {
            'place_id': place_id,
            'fields': 'name,place_id,url,type,opening_hours,utc_offset,reservable,rating,price_level,review,user_ratings_total,formatted_phone_number,international_phone_number,website,scope,vicinity,formatted_address,business_status,dine_in,editorial_summary,serves_breakfast,serves_brunch,serves_dinner,serves_lunch,serves_wine,serves_vegetarian_food,takeout',
            'key': self.api_key
        }
        response = requests.get(details_url, params=params)
        if response.status_code == 200:
            details_data = response.json()
            if details_data['status'] == 'OK':
                return details_data['result']
        return {}

    def save_research_report(self, filename):
        with open(filename, 'w') as file:
            json.dump(self.data, file, indent=4)

    ### WEB SCRAPING METHODS TO APPEND MORE DATA TO THE REPORT DATASET ###

    def scrape_website_data(self, url):
        print(f"Scraping website: {url}")
        try:
            response = requests.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')

            # Base URL extraction for concatenating with relative links
            base_url = '{uri.scheme}://{uri.netloc}'.format(uri=urlparse(url))

            # Extract and filter links
            links = soup.find_all('a', href=True)
            menu_links = set(self.get_full_link(link['href'], base_url) for link in links if 'menu' in link.get_text().lower() and self.is_valid_link(link['href']))
            event_links = set(self.get_full_link(link['href'], base_url) for link in links if any(word in link.get_text().lower() for word in ['event', 'calendar', 'schedule']) and self.is_valid_link(link['href']))
            pdf_links = set(self.get_full_link(link['href'], base_url) for link in links if '.pdf' in link['href'] and self.is_valid_link(link['href']))

            return {'menu_links': list(menu_links), 'event_links': list(event_links), 'pdf_links': list(pdf_links)}

        except requests.exceptions.Timeout:
            print(f"Timeout occurred while scraping {url}")
            return {}
        except requests.exceptions.RequestException as e:
            print(f"Error scraping {url}: {e}")
            return {}

    def get_full_link(self, link, base_url):
        if link.startswith('/'):
            return urljoin(base_url, link)
        return link

    def is_valid_link(self, link):
        # Add logic to filter out invalid or duplicate links
        return link.startswith("http") or (link.startswith("/") and not link.startswith("//"))
        
    def generate_summary(self):
        for group, places in self.data['grouped_places'].items():
            for place in tqdm(places, desc=f"Generating summaries for {group}"):
                try:
                    place_data_str = json.dumps(place, ensure_ascii=False)
                    prompt = (
                        "Hi Gemini, you are being asked to write a *CONCISE AND SHORT* summary - please examine the following json object which contains data about a restaurant. "
                        "The restaurant in this json object is within physical proximity of the user's American fine dining restaurant and the user needs *competitive business insights* about the restaurant in this json object. "
                        "The user's restaurant is a luxury American fine dining establishment with a robust wine list and elevated style of service. "
                        "This analysis is to *understand the nuances* of this available business data and augment the json onject with your observations for the user. The user is an operations manager who is analyzing businesses in a potential upcoming market near their restaurant. "
                        "Your observations will be added as a new 'gemini_summary' field into the current json object. "
                        "You are writing a short and concise summary of the business in this json object for the user to better understand how this business will affect and compete with their own. "
                        "You are being asked to compare the competitive relevance and business related relevance between the business in this object and the user's American fine dining restaurant. "
                        "The user's restaurant is an American fine dining restaurant, located near the restaurant represented in in this json object. "
                        "The distance between the restaurant in the json object and the user's restaurant is listed as a field in the json object. "
                        "The user needs to examine potential competitors in their area and draw business insights for their own strategy. "
                        "Please write a summary for the pre-planning specialistists and operations managers of this American fine dining restaurant, helping them understand this business in the context of their relative market. "
                        "Tell them everything they might want to know about this potential competitor in their market and how it related to their American fine dining restaurant located nearby. "
                        "Look at the business status, editorial summary, and website for this business listed in the json data and understand the business through the eyes of the operations managers and pre-planning specialists of the American fine dining restaurant. "
                        "Also analyze the overall sentiment tone and customer messages in the reviews for this business. "
                        "Include a S.W.O.T. analysis of the business in this json object as it relates to the user's competing American fine dining restaurant located nearby. "
                        "Also include any of your own knowledge about the business in this json object, and its relevance to the user's restaurant business in the same market. "
                        "Your output should be a short and concise summary of the business in this json object. "
                        "Please keep it simple and concise. Statements like 'this business is not relevant to you because the style of food and price tier are much different' or 'this business is relevant to you because they have wine, they accept reservations, they are expensive, and they are luxury American food' are acceptable. "
                        "Please limit your responses to a few sentences. "
                        "Do not include any markdown or special characters in your final output. "
                        "Please limit your responses to no more than 1-2 paragraphs of direct and concise text. "
                        "Be consice and direct. The user has to read many of these summaries and they need to be able to quickly understand the business in this json object. "
                        f"Here is the object to analyze: {place_data_str}"
                    )
                    response = model.generate_content(prompt)
                    summary = response.text
                    place['gemini_summary'] = summary
                    time.sleep(1)  # Pause for 1 second between API calls
                except Exception as e:
                    print(f"Error generating summary for {place.get('name', 'N/A')}: {e}")
                    place['gemini_summary'] = 'N/A'

        # Optionally return the whole data if needed
        return self.data

    def augment_place_details_with_web_data(self):
        print("Augmenting place details with web data")
        for group, places in self.data['grouped_places'].items():
            for place in tqdm(places, desc=f"Processing {group}"):
                if 'website' in place:
                    web_data = self.scrape_website_data(place['website'])
                    place.update(web_data)
                    
        self.generate_summary()
        
    ### PRINTING THE FINDINGS TO THE CONSOLE ###
    
    def print_human_readable(self):
        print(f"Research Report for Address: {self.data['address']}\n")
        for group, places in self.data['grouped_places'].items():
            print(f"{group.capitalize()}:")
            for place in places:
                distance = place.get('distance_from_target_address', 'N/A')
                print(f"Distance from Address: {distance} km")
                self.print_place_details(place)
            print("\n")
    
    def print_place_details(self, place):
        # Print basic information
        print(f"Name: {place.get('name', 'N/A')}")
        print(f"Types: {', '.join(place.get('types', ['N/A']))}")
        print(f"Rating: {place.get('rating', 'N/A')}")
        print(f"Price Level: {'$' * place.get('price_level', 0)}")
        
        # Iterate through additional fields
        additional_fields = [
            'formatted_phone_number', 'user_ratings_total', 'website', 'vicinity',
            'formatted_address', 'business_status', 'dine_in', 'editorial_summary',
            'serves_breakfast', 'serves_brunch', 'serves_dinner', 'serves_lunch',
            'serves_wine', 'serves_vegetarian_food', 'takeout'
        ]
        for field in additional_fields:
            value = place.get(field, 'N/A')
            if isinstance(value, list):
                value = ', '.join(value)
            print(f"{field.replace('_', ' ').capitalize()}: {value}")

        # Print links
        print("Menu Links:", ", ".join(place.get('menu_links', [])))
        print("Event Links:", ", ".join(place.get('event_links', [])))
        print("PDF Links:", ", ".join(place.get('pdf_links', [])))
        print("AI Summary:")
        print(place.get('gemini_summary', 'N/A'))
        print("-------------------------------------------------------")
        print()
        print()
    
    ### FILE SAVING METHODS FOR DOCX AND CSV ###

    def save_report_as_word(self, filename):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Open Sans'
        font.size = Pt(12)

        doc.add_heading('Research Report', level=1)
        doc.add_paragraph(f'Address: {self.data["address"]}\n')
        
        self.add_summary_to_doc(doc)

        for group, places in self.data['grouped_places'].items():
            doc.add_heading(f'{group.capitalize()}:', level=2)
            for place in places:
                if isinstance(place, dict):
                    self.add_place_details_to_doc(doc, place)
                else:
                    print(f"Warning: Expected a dictionary for 'place', got {type(place)}")

        doc.save(filename)

    def add_place_details_to_doc(self, doc, place):
        # Ensure 'place' is a dictionary
        if not isinstance(place, dict):
            print(f"Warning: Expected a dictionary for 'place', got {type(place)}")
            return

        doc.add_heading(place.get('name', 'N/A'), level=3)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Data'
    
        # Standard fields
        standard_fields = ['editorial_summary', 'business_status', 'website', 'url', 'types', 'price_level', 'formatted_phone_number', 'international_phone_number', 'formatted_address', 'rating', 'user_ratings_total', 'review']
        for field in standard_fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field.replace('_', ' ').capitalize()
            row_cells[1].text = str(place.get(field, 'N/A'))

        # Additional fields
        additional_fields = ['dine_in', 'takeout', 'reservable', 'serves_breakfast', 'serves_brunch', 'serves_lunch', 'serves_dinner', 'serves_wine', 'serves_vegetarian_food']
        for field in additional_fields:
            value = place.get(field, 'N/A')
            value_text = 'Yes' if value is True else 'No' if value is False else str(value)
            row_cells = table.add_row().cells
            row_cells[0].text = field.replace('_', ' ').capitalize()
            row_cells[1].text = value_text

        # List fields
        self.add_list_to_doc(doc, 'Menu Links', place.get('menu_links', []))
        self.add_list_to_doc(doc, 'Event Links', place.get('event_links', []))
        self.add_list_to_doc(doc, 'PDF Links', place.get('pdf_links', []))
        
        doc.add_paragraph(f"AI Summary: {place.get('gemini_summary', 'N/A')}", style='List Bullet')


    def add_list_to_doc(self, doc, title, items):
        if items:
            doc.add_paragraph(f'{title}:', style='List Bullet')
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

    def add_summary_to_doc(self, doc):
        doc.add_heading('Summary of Findings', level=2)
        for group, places in self.data['grouped_places'].items():
            doc.add_heading(f'{group.capitalize()}:', level=3)
            
            # Create a table for this group with an additional column for the weighted score
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Weighted Match Score'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Rating'
            hdr_cells[3].text = 'Price Level'
            hdr_cells[4].text = 'Distance (km)'

            for place in places:
                row_cells = table.add_row().cells
                row_cells[0].text = str(round(place.get('weighted_score', 0), 2))  # Weighted score rounded to 2 decimal places
                row_cells[1].text = place.get('name', 'N/A')
                row_cells[2].text = str(place.get('rating', 'N/A'))
                row_cells[3].text = "$" * place.get('price_level', 0)
                row_cells[4].text = f"{round(place.get('distance_from_target_address', 0), 1)}"

    def save_report_as_csv(self, filename):
        data_for_df = []
        for group, places in self.data['grouped_places'].items():
            for place in places:
                place_data = {'Name': place.get('name', 'N/A'), 
                              'Type': ', '.join(place.get('types', ['N/A'])),
                              'Gemini Summary': place.get('gemini_summary', 'N/A'),
                              }
                fields = [
                    'editorial_summary', 'business_status', 'website', 'url', 'types', 
                    'price_level', 'formatted_phone_number', 'international_phone_number', 
                    'formatted_address', 'rating', 'user_ratings_total', 'dine_in', 
                    'takeout', 'reservable', 'serves_breakfast', 'serves_brunch', 
                    'serves_lunch', 'serves_dinner', 'serves_wine', 'serves_vegetarian_food'
                ]
                place_data.update({field: place.get(field, 'N/A') for field in fields})
                
                # List fields
                place_data['Menu Links'] = ', '.join(place.get('menu_links', []))
                place_data['Event Links'] = ', '.join(place.get('event_links', []))
                place_data['PDF Links'] = ', '.join(place.get('pdf_links', []))

                # Distance and score
                place_data['Distance from Target Address (km)'] = place.get('distance_from_target_address', 'N/A')
                place_data['Weighted Score'] = self.calculate_weighted_score(place)

                data_for_df.append(place_data)

        df = pd.DataFrame(data_for_df)
        df.to_csv(filename, index=False)

if __name__ == '__main__':
    address_researcher = AddressResearcher(google_maps_api_key)
    address_researcher.perform_research(address)
    address_researcher.augment_place_details_with_web_data()
    address_researcher.print_human_readable()
    address_researcher.save_research_report(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented.json")
    address_researcher.save_report_as_word(f"{FILE_DROP_DIR_PATH}/research_report_doc_formatted.docx")
    address_researcher.save_report_as_csv(f"{FILE_DROP_DIR_PATH}/research_report_spreadsheet.csv")

