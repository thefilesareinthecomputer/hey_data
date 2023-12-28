
# IMPORTS ###################################################################################################################################

from bs4 import BeautifulSoup
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from fake_useragent import UserAgent
from math import radians, cos, sin, asin, sqrt
from openai import OpenAI
from pyppeteer import launch, errors as pyppeteer_errors
from tqdm import tqdm
from transformers import MarianMTModel, MarianTokenizer
from urllib.parse import urlparse, urljoin
import asyncio # new
import certifi
import datetime
import google.generativeai as genai
import json
import logging
import numpy as np
import os
import pandas as pd
import pyautogui
import pytz
import queue
import random
import re
import requests
import speech_recognition as sr
import ssl
import subprocess
import threading
import time
import tkinter as tk
import traceback
import webbrowser
import wikipedia
import wolframalpha
import yfinance as yf

# CUSTOM IMPORTS ##############################################################################################################################


# CONSTANTS ###################################################################################################################################

# bring in the environment variables from the .env file
load_dotenv()
ACTIVATION_WORD = os.getenv('ACTIVATION_WORD', 'robot')
USER_DOWNLOADS_FOLDER = os.getenv('USER_DOWNLOADS_FOLDER')
USER_PREFERRED_LANGUAGE = os.getenv('USER_PREFERRED_LANGUAGE', 'en')  # 2-letter lowercase
USER_PREFERRED_VOICE = os.getenv('USER_PREFERRED_VOICE', 'Daniel')  # Daniel
USER_PREFERRED_NAME = os.getenv('USER_PREFERRED_NAME', 'User')  # Title case
USER_SELECTED_PASSWORD = os.getenv('USER_SELECTED_PASSWORD', 'None')  
USER_SELECTED_HOME_CITY = os.getenv('USER_SELECTED_HOME_CITY', 'None')  # Title case
USER_SELECTED_HOME_COUNTY = os.getenv('USER_SELECTED_HOME_COUNTY', 'None')  # Title case
USER_SELECTED_HOME_STATE = os.getenv('USER_SELECTED_HOME_STATE', 'None')  # Title case
USER_SELECTED_HOME_COUNTRY = os.getenv('USER_SELECTED_HOME_COUNTRY', 'None')  # 2-letter country code
USER_SELECTED_HOME_LAT = os.getenv('USER_SELECTED_HOME_LAT', 'None')  # Float with 6 decimal places
USER_SELECTED_HOME_LON = os.getenv('USER_SELECTED_HOME_LON', 'None')  # Float with 6 decimal places 
USER_SELECTED_TIMEZONE = os.getenv('USER_SELECTED_TIMEZONE', 'America/Chicago')  # Country/State format
USER_STOCK_WATCH_LIST = os.getenv('USER_STOCK_WATCH_LIST', 'None').split(',')  # Comma separated list of stock symbols
PROJECT_VENV_DIRECTORY = os.getenv('PROJECT_VENV_DIRECTORY')

# establish relative file paths for the current script
SCRIPT_DIR_PATH = os.path.dirname(os.path.realpath(__file__))
PROJECT_ROOT_DIR_PATH = os.path.dirname(SCRIPT_DIR_PATH)
ARCHIVED_DEV_VERSIONS_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, '_archive')
FILE_DROP_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_generated_files')
LOCAL_LLMS_DIR = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_local_models')
NOTES_DROP_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_base_knowledge')
SOURCE_DATA_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, 'app_source_data')
TESTS_DIR_PATH = os.path.join(PROJECT_ROOT_DIR_PATH, '_tests')
UTILITIES_DIR_PATH = os.path.join(SCRIPT_DIR_PATH, 'utilities')

folders_to_create = [ARCHIVED_DEV_VERSIONS_PATH, FILE_DROP_DIR_PATH, LOCAL_LLMS_DIR, NOTES_DROP_DIR_PATH, SOURCE_DATA_DIR_PATH, TESTS_DIR_PATH, UTILITIES_DIR_PATH]
for folder in folders_to_create:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Set API keys and other information from environment variables
open_weather_api_key = os.getenv('OPEN_WEATHER_API_KEY')
wolfram_app_id = os.getenv('WOLFRAM_APP_ID')
openai_api_key=os.getenv('OPENAI_API_KEY')
google_cloud_api_key = os.getenv('GOOGLE_CLOUD_API_KEY')
google_maps_api_key = os.getenv('GOOGLE_MAPS_API_KEY')
google_gemini_api_key = os.getenv('GOOGLE_GEMINI_API_KEY')

# Initialize LLM
genai.configure(api_key=google_gemini_api_key)
model = genai.GenerativeModel('gemini-pro')

# Set the terminal output display options
pd.set_option('display.max_columns', 10)
pd.set_option('display.max_rows', 150)
pd.set_option('display.width', None)
pd.set_option('display.max_colwidth', 35)
pd.set_option('display.expand_frame_repr', False)
pd.set_option('display.precision', 1) 

# FUNCTION DEFINITIONS ###################################################################################################################################

address = "9 9th Ave, New York, NY 10014"
search_distance = 10000
    
class AddressResearcher:
    def __init__(self, api_key):
        self.api_key = api_key
        self.geocode_url = "https://maps.googleapis.com/maps/api/geocode/json"
        self.place_search_url = "https://maps.googleapis.com/maps/api/place/nearbysearch/json"
        self.data = {}
        self.type_groupings = {
            'restaurants': ['restaurant'],
        }

    # User-Agent List
    USER_AGENTS = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.0.3 Safari/605.1.15',
        # ... add more user agents ...
    ]
    
    ### INITIAL PASS OF REPORT FROM API CALL ###

    def perform_research(self, address):
        distance = search_distance
        print(f"Starting research for address: {address}")
        lat, lng = self.geocode_address(address)
        if lat is None or lng is None:
            print("Geocoding failed.")
            return

        all_places = self.search_nearby(lat, lng, distance)

        # Grouping and sorting
        grouped_places = {}
        for place in all_places:
            group = place['group']
            place['weighted_score'] = self.calculate_weighted_score(place)
            if group not in grouped_places:
                grouped_places[group] = []
            grouped_places[group].append(place)

        for group in grouped_places:
            grouped_places[group].sort(key=lambda x: x['weighted_score'], reverse=True)

        self.data = {'address': address, 'grouped_places': grouped_places}
    
    def geocode_address(self, address):
        params = {'address': address, 'key': self.api_key}
        response = requests.get(self.geocode_url, params=params)
        if response.status_code != 200:
            print(f"Geocode API request failed with status code: {response.status_code}")
            return None, None
        geocode_data = response.json()
        if geocode_data['status'] != 'OK':
            print(f"Geocode API response error: {geocode_data['status']}")
            return None, None
        location = geocode_data['results'][0]['geometry']['location']
        return location['lat'], location['lng']

    def haversine(self, lon1, lat1, lon2, lat2):
        """
        Calculate the great circle distance in kilometers between two points 
        on the earth (specified in decimal degrees)
        """
        # convert decimal degrees to radians 
        lon1, lat1, lon2, lat2 = map(radians, [lon1, lat1, lon2, lat2])

        # Haversine formula 
        dlon = lon2 - lon1 
        dlat = lat2 - lat1 
        a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
        c = 2 * asin(sqrt(a)) 
        km = 6371 * c # Radius of earth in kilometers
        return km

    def search_nearby(self, lat, lng, radius):
        all_places = []
        missing_coordinate_places = []

        for group, types in self.type_groupings.items():
            for place_type in types:
                next_page_token = None
                while True:
                    params = {
                        'location': f"{lat},{lng}",
                        'radius': radius,
                        'type': place_type,
                        'key': self.api_key,
                        'pagetoken': next_page_token
                    }
                    response = requests.get(self.place_search_url, params=params)
                    if response.status_code == 200:
                        search_data = response.json()
                        if search_data['status'] == 'OK':
                            for place in search_data['results']:
                                place_details = self.get_place_details(place['place_id'])
                                if place_details:
                                    place_lat = place_details.get('geometry', {}).get('location', {}).get('lat')
                                    place_lng = place_details.get('geometry', {}).get('location', {}).get('lng')
                                    if place_lat and place_lng:
                                        distance = self.haversine(lng, lat, place_lng, place_lat)
                                        place_details['distance_from_target_address'] = distance
                                    else:
                                        missing_coordinate_places.append(place_details)
                                    place_details['group'] = group
                                    all_places.append(place_details)
                            next_page_token = search_data.get('next_page_token')
                            if not next_page_token:
                                break
                            time.sleep(2)  # Delay for token validity
                        else:
                            break
                    else:
                        break

        # Process the batch of missing coordinate places
        self.process_missing_coordinate_places(missing_coordinate_places, lat, lng)

        return all_places

    def process_missing_coordinate_places(self, missing_places, origin_lat, origin_lng):
        # Google Maps Distance Matrix API endpoint
        distance_matrix_url = "https://maps.googleapis.com/maps/api/distancematrix/json"

        for place in missing_places:
            # Building the address from the place details
            destination_address = place.get('vicinity', 'Unknown Location')
            params = {
                'origins': f"{origin_lat},{origin_lng}",
                'destinations': destination_address,
                'key': self.api_key
            }
            response = requests.get(distance_matrix_url, params=params)
            if response.status_code == 200:
                matrix_data = response.json()
                if matrix_data['status'] == 'OK':
                    # Extracting distance information
                    result = matrix_data['rows'][0]['elements'][0]
                    if result['status'] == 'OK':
                        distance_value = result['distance']['value'] / 1000  # Convert meters to kilometers
                        place['distance_from_target_address'] = distance_value
                    else:
                        print(f"Distance calculation failed for {destination_address}: {result['status']}")
                else:
                    print(f"Distance Matrix API error: {matrix_data['status']}")
            else:
                print(f"Distance Matrix API request failed with status code: {response.status_code}")
    
    def calculate_weighted_score(self, place):
        rating = place.get('rating', 0)
        price_level = place.get('price_level', 0)
        distance = place.get('distance_from_target_address', float('inf'))

        # Adjust weights as needed
        # Higher weight for price level and proximity, excluding user ratings total
        distance_factor = 1 / (distance + 1)  # Add 1 to avoid division by zero
        return (rating * 0.4 + price_level * 0.6) * distance_factor

    def get_place_details(self, place_id):
        print(f"Fetching details for place ID: {place_id}")
        details_url = "https://maps.googleapis.com/maps/api/place/details/json"
        params = {
            'place_id': place_id,
            'fields': 'name,place_id,url,type,opening_hours,utc_offset,reservable,rating,price_level,review,user_ratings_total,formatted_phone_number,international_phone_number,website,scope,vicinity,formatted_address,business_status,dine_in,editorial_summary,serves_breakfast,serves_brunch,serves_dinner,serves_lunch,serves_wine,serves_vegetarian_food,takeout',
            'key': self.api_key
        }
        response = requests.get(details_url, params=params)
        if response.status_code == 200:
            details_data = response.json()
            if details_data['status'] == 'OK':
                return details_data['result']
        return {}

    def save_research_report(self, filename):
        with open(filename, 'w') as file:
            json.dump(self.data, file, indent=4)

    ### WEB SCRAPING METHODS TO APPEND MORE DATA TO THE REPORT DATASET ###
        
    def generate_summary(self):
        for group, places in self.data['grouped_places'].items():
            for place in tqdm(places, desc=f"Generating summaries for {group}"):
                try:
                    place_data_str = json.dumps(place, ensure_ascii=False)
                    prompt = (
                        "Hi Gemini, write a *CONCISE AND SHORT* summary. Please begin by examining the following json object. This json contains data about a business. "
                        "The business in this json object is located near an American fine dining restaurant. The goal is to gain *market business insights* about the business in this json object. "
                        "The nearby comparison restaurant is a luxury American fine dining establishment with a robust wine list and elevated style of service. "
                        "The goal is to *understand the nuances* of this json data and append the json object with your observations. The user is a research analyst who is researching the affect of businesses located near a restaurant. "
                        "You are writing a short and concise summary of the business in this json object and how this business will affect the American fine dining restaurant either positively or negatively. "
                        "You are being asked to compare the competitive business related relevance between the business in this object and the American fine dining restaurant. We want to learn if it will help or hinder the nearby restaurant. "
                        "The comparison restaurant is an American fine dining restaurant, located near the business in the json object. "
                        "The distance between the business in the json object and the comparison restaurant is listed in the json object. "
                        "The goal is to examine potential competitors and/or partners in the area of the American fine dining restaurant and draw business hypotheses about their potential relationship. "
                        "Please write a summary of how this business could affect a nearby American fine dining restaurant. "
                        "Explain everything about this business as a potential competitor or partner within the American restaurant's local market. "
                        "Also analyze the overall sentiment tone and customer messages in the reviews. "
                        "Include a S.W.O.T. analysis of the business in this json object as it relates to the comparison American fine dining restaurant located nearby. "
                        "Only analyze how the nearby business may passivle influence the nearby restaurant. Do not make any recommendations about partnering directly. "
                        "Also include any of your own knowledge about the listed business and its relevance to the restaurant located nearby in the same market. "
                        "Your output should be a short and concise summary of the business in this json object and the details outlined in this message. "
                        "Please keep it simple and concise. Statements like 'this business is not relevant to you because the style of food and price tier are much different' or 'this business is relevant to you because they have wine, they accept reservations, they are expensive, and they are luxury American food' are acceptable. "
                        "Please limit your responses to a few sentences. "
                        "Do not include any markdown or special characters in your final output. "
                        "Please limit your responses to no more than 1-2 paragraphs of direct and concise text. "
                        "Be consice and direct. "
                        f"Here is the object to analyze: {place_data_str}"
                    )
                    response = model.generate_content(prompt)
                    summary = response.text
                    place['gemini_summary'] = summary
                    time.sleep(2) 
                except Exception as e:
                    print(f"Error generating summary for {place.get('name', 'N/A')}: {e}")
                    place['gemini_summary'] = 'N/A'
                    time.sleep(2)

        # Optionally return the whole data if needed
        return self.data

    def business_summary(self):
        for group, places in self.data['grouped_places'].items():
            for place in tqdm(places, desc=f"Generating summaries for {group}"):
                try:
                    place_data_str = json.dumps(place, ensure_ascii=False)
                    prompt = (
                        "Hi Gemini, write a *CONCISE AND SHORT* summary. Please begin by examining the following json object. This json contains data about a business. "
                        "Obtain the info about the business from the JSON like the business name and location. "
                        "Write a summary that describes all *BUSINESS CRITICAL* information about the business in the json object. "
                        "The goal is to *understand the nuances* of this json data and append the json object with your observations. The user is a research analyst who is researching the business and its market presence and overall company info. "
                        "The goal is to examine this business and its role within the context of its local market. "
                        "Explain everything you know from your training data about this business at this specific location. "
                        "List all important headlines and national news updates about this location (if any). "
                        "List any awards or accolades this business has recieved such as Michelin stars, James Beard awards, etc. "
                        "Provide any known information about the business such as whether they have any Michelin stars, have been in the news, if they have a famous chef, etc. "
                        "Also include any of your own knowledge about the listed business. "
                        "Your output should be a short and concise summary of the business, with your own insights included in the summary. "
                        "Please keep it simple and concise and add value to the knowledge with factual data. "
                        "Please limit your responses to a few sentences, but include a lot of information. "
                        "Do not include any markdown or special characters in your final output. "
                        "Please limit your responses to no more than 1-2 paragraphs of direct and concise text. "
                        "Be consice and direct and substantial and thoughtful and critical and detail-oriented and intelligent and thorough. "
                        f"Here is the object to analyze: {place_data_str}"
                    )
                    response = model.generate_content(prompt)
                    summary = response.text
                    place['business_summary'] = summary
                    time.sleep(2) 
                except Exception as e:
                    print(f"Error generating summary for {place.get('name', 'N/A')}: {e}")
                    place['business_summary'] = 'N/A'
                    time.sleep(2)

        # Optionally return the whole data if needed
        return self.data

    # def product_summary(self):
    #     # Initialize LLM
    #     genai.configure(api_key=google_gemini_api_key)
    #     model = genai.GenerativeModel('gemini-pro')
    
    #     for group, places in self.data['grouped_places'].items():
    #         for place in tqdm(places, desc=f"Generating summaries for {group}"):
    #             try:
    #                 place_data_str = json.dumps(place, ensure_ascii=False)
    #                 prompt = (
    #                     "Hi Gemini, write a *CONCISE AND SHORT* summary. Please begin by examining the following json object. This json contains data about a business. "
    #                     "Describe all products available for sale at this business. "
    #                     f"Here is the object to analyze: {place_data_str}"
    #                 )
    #                 response = model.generate_content(prompt)
    #                 summary = response.text
    #                 place['business_summary'] = summary
    #                 time.sleep(2) 
    #             except Exception as e:
    #                 print(f"Error generating summary for {place.get('name', 'N/A')}: {e}")
    #                 place['product_summary'] = 'N/A'
    #                 time.sleep(2)

    #     # Optionally return the whole data if needed
    #     return self.data
    
    # def product_summary(self):
    #     for group, places in self.data['grouped_places'].items():
    #         for place in tqdm(places, desc=f"Generating summaries for {group}"):
    #             try:
    #                 # Combining JSON data with scraped menu content
    #                 menu_content = place.get('menu_content', '')
    #                 place_data_str = json.dumps(place, ensure_ascii=False)
    #                 prompt = (
    #                     "Hi Gemini, write a *CONCISE AND SHORT* summary. This json contains data about a business, "
    #                     "including a description of their menu. Describe all products available for sale at this business. "
    #                     f"Here is the object to analyze: {place_data_str}\n\nMenu/Products Description:\n{menu_content}"
    #                 )
    #                 response = model.generate_content(prompt)
    #                 summary = response.text
    #                 place['product_summary'] = summary
    #                 time.sleep(2)
    #             except Exception as e:
    #                 print(f"Error generating summary for {place.get('name', 'N/A')}: {e}")
    #                 place['product_summary'] = 'N/A'
    #                 time.sleep(2)

    #     # Optionally return the whole data if needed
    #     return self.data

    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////

    
    # async def scrape_website_data(self, browser, url, visited_urls=set(), depth=0, max_depth=3):
    #     if depth > max_depth or url in visited_urls:
    #         return None

    #     visited_urls.add(url)

    #     try:
    #         print(f"Scraping website: {url} at depth {depth}")
    #         page = await browser.newPage()
    #         await page.setUserAgent(UserAgent().random)
    #         response = await page.goto(url, {'waitUntil': 'networkidle2', 'timeout': 60000})

    #         if response and not response.ok:  # Check if response is successful
    #             print(f"Error loading {url}: {response.status}")
    #             await page.close()
    #             return None

    #         if depth < max_depth:
    #             links = await page.evaluate('''() => {
    #                 return Array.from(document.querySelectorAll('a')).map(a => a.href);
    #             }''')

    #             for link in links:
    #                 if self.is_valid_link(link) and self.is_potential_menu_link(link):
    #                     content = await self.scrape_website_data(browser, link, visited_urls, depth + 1, max_depth)
    #                     if content:
    #                         await page.close()
    #                         return content

    #         if await page.isClosed():  # Check if the page is still active
    #             return None

    #         current_page_content = await page.content()
    #         soup = BeautifulSoup(current_page_content, 'html.parser')
    #         menu_content = soup.get_text()
    #         await page.close()
    #         return {'menu_link': url, 'menu_content': menu_content.strip()}

    #     except Exception as e:
    #         print(f"Error scraping {url}: {str(e)}")
    #         return None

    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////

    # async def scrape_website_data(self, browser, url, depth=0, max_depth=3, visited_urls=None):
    #     # Convert depth to an integer to prevent type errors
    #     depth = int(depth)

    #     if depth > max_depth:
    #         return None

    #     if visited_urls is None:
    #         visited_urls = set()

    #     if url in visited_urls:
    #         return None

    #     visited_urls.add(url)

    #     try:
    #         browser = await launch(headless=False, args=['--no-sandbox', '--disable-setuid-sandbox'])
    #         page = await browser.newPage()
    #         await page.setUserAgent(UserAgent().random)
    #         await page.goto(url, {'waitUntil': 'networkidle2'})

    #         links = await page.evaluate('''() => Array.from(document.querySelectorAll('a')).map(a => a.href);''')

    #         for link in links:
    #             if self.is_valid_link(link) and self.is_potential_menu_link(link) and link not in visited_urls:
    #                 content = await self.scrape_website_data(link, depth + 1, max_depth, visited_urls)
    #                 if content:
    #                     await browser.close()
    #                     return content

    #         menu_content = await page.evaluate('document.body.innerText')
    #         await browser.close()
    #         return {'menu_link': url, 'menu_content': menu_content.strip()}

    #     except Exception as e:
    #         print(f"Error scraping {url}: {str(e)}")
    #         return None

    #     finally:
    #         await browser.close()

    ####################################//////////////////////////////////////////////////////////////////////////////////////////////////////////
  
    # async def scrape_website_data(self, browser, url, visited_urls=set(), depth=0, max_depth=2):
    #     """
    #     Recursively scrape website data.
    #     :param browser: Pyppeteer browser instance.
    #     :param url: URL to scrape.
    #     :param visited_urls: Set of already visited URLs to avoid repetition.
    #     :param depth: Current depth of recursion.
    #     :param max_depth: Maximum depth allowed for recursion.
    #     :return: Dictionary with menu link and content if found.
    #     """
    #     if depth > max_depth or url in visited_urls:
    #         return None

    #     visited_urls.add(url)

    #     try:
    #         page = await browser.newPage()
    #         # User-Agent Rotation
    #         user_agent = random.choice(self.USER_AGENTS)
    #         await page.setUserAgent(user_agent)

    #         # Set Referrer
    #         await page.setExtraHTTPHeaders({'Referer': 'https://www.google.com/'})
            
    #         await page.goto(url, {'waitUntil': 'networkidle2'})
    #         # Throttling Requests
    #         await asyncio.sleep(random.uniform(0.5, 2.0))

    #         # Extract all links from the page
    #         links = await page.evaluate('''() => {
    #             return Array.from(document.querySelectorAll('a')).map(a => a.href);
    #         }''')

    #         # Recursively check each link
    #         for link in links:
    #             if self.is_valid_link(link) and self.is_potential_menu_link(link):
    #                 content = await self.scrape_website_data(browser, link, visited_urls, depth + 1, max_depth)
    #                 if content:
    #                     await page.close()
    #                     return content

    #         # Extract content from the current page
    #         current_page_content = await page.content()
    #         soup = BeautifulSoup(current_page_content, 'html.parser')
    #         menu_content = soup.get_text()

    #         await page.close()
    #         return {'menu_link': url, 'menu_content': menu_content.strip()}

    #     except Exception as e:
    #         print(f"Error scraping {url}: {str(e)}")
    #         return None
        
    # async def scrape_website_data(self, url, depth=0, max_depth=12):
    #     depth = int(depth)
    #     if depth > max_depth:
    #         return None

    #     try:
    #         browser = await launch(headless=False, args=['--no-sandbox', '--disable-setuid-sandbox'])
    #         page = await browser.newPage()
    #         await page.setUserAgent(UserAgent().random)
    #         await page.setExtraHTTPHeaders({'Referer': 'https://www.google.com'})
            
    #         try:
    #             await page.goto(url, {'waitUntil': 'networkidle2'})
    #         except pyppeteer_errors.PageError as e:
    #             logging.error(f"Page navigation failed for {url}: {e}")
    #             return None

    #         # Manual intervention for CAPTCHAs
    #         await asyncio.sleep(30)

    #         links = await page.evaluate('''() => Array.from(document.querySelectorAll('a')).map(a => a.href);''')

    #         for link in links:
    #             if self.is_valid_link(link) and self.is_potential_menu_link(link):
    #                 content = await self.scrape_website_data(link, depth + 1, max_depth)
    #                 if content:
    #                     return content

    #         content = await page.content()
    #         soup = BeautifulSoup(content, 'html.parser')
    #         menu_content = soup.get_text()

    #         return {'menu_link': url, 'menu_content': menu_content.strip()}

    #     except Exception as e:
    #         logging.error(f"Error scraping {url} at depth {depth}: {e}")
    #         return None

    #     finally:
    #         if browser:
    #             await browser.close()

    # async def scrape_website_data(self, browser, url, visited_urls=set(), depth=0, max_depth=3):
    #     if depth > max_depth or url in visited_urls:
    #         return None

    #     visited_urls.add(url)

    #     try:
    #         print(f"Scraping website: {url} at depth {depth}")
    #         page = await browser.newPage()
    #         await page.setUserAgent(UserAgent().random)
    #         response = await page.goto(url, {'waitUntil': 'networkidle2', 'timeout': 60000})

    #         if response and not response.ok:  # Check if response is successful
    #             print(f"Error loading {url}: {response.status}")
    #             await page.close()
    #             return None

    #         if depth < max_depth:
    #             links = await page.evaluate('''() => {
    #                 return Array.from(document.querySelectorAll('a')).map(a => a.href);
    #             }''')

    #             for link in links:
    #                 if self.is_valid_link(link) and self.is_potential_menu_link(link):
    #                     content = await self.scrape_website_data(browser, link, visited_urls, depth + 1, max_depth)
    #                     if content:
    #                         await page.close()
    #                         return content

    #         if await page.isClosed():  # Check if the page is still active
    #             return None

    #         current_page_content = await page.content()
    #         soup = BeautifulSoup(current_page_content, 'html.parser')
    #         menu_content = soup.get_text()
    #         await page.close()
    #         return {'menu_link': url, 'menu_content': menu_content.strip()}

    #     except Exception as e:
    #         print(f"Error scraping {url}: {str(e)}")
    #         return None

    async def scrape_website_data(self, browser, url, depth=0, max_depth=3, visited_urls=None):
        # Convert depth to an integer to prevent type errors
        depth = int(depth)

        if depth > max_depth:
            return None

        if visited_urls is None:
            visited_urls = set()

        if url in visited_urls:
            return None

        visited_urls.add(url)

        try:
            browser = await launch(headless=False, args=['--no-sandbox', '--disable-setuid-sandbox'])
            page = await browser.newPage()
            await page.setUserAgent(UserAgent().random)
            await page.goto(url, {'waitUntil': 'networkidle2'})

            links = await page.evaluate('''() => Array.from(document.querySelectorAll('a')).map(a => a.href);''')

            for link in links:
                if self.is_valid_link(link) and self.is_potential_menu_link(link) and link not in visited_urls:
                    content = await self.scrape_website_data(link, depth + 1, max_depth, visited_urls)
                    if content:
                        await browser.close()
                        return content

            menu_content = await page.evaluate('document.body.innerText')
            await browser.close()
            return {'menu_link': url, 'menu_content': menu_content.strip()}

        except Exception as e:
            print(f"Error scraping {url}: {str(e)}")
            return None

        finally:
            await browser.close()
        
    def is_potential_menu_link(self, url):
        """
        Check if a URL is a potential menu or product link.
        :param url: URL to check.
        :return: Boolean indicating if the URL is a potential menu/product link.
        """
        menu_keywords = [
            'menu', 'products', 'dinner', 'food', 'dishes', 'cuisine',
            'dining', 'lunch', 'breakfast', 'brunch', 'specials', 'offers',
            'meals', 'courses', 'drinks', 'beverages', 'eatery', 'gastronomy',
            'culinary', 'delicacies', 'fare', 'sustenance', 'provisions'
        ]
        return any(keyword in url.lower() for keyword in menu_keywords)

    # def sync_scrape_website_data(self, browser, url):
    #     return asyncio.get_event_loop().run_until_complete(self.scrape_website_data(browser, url))
    
    def sync_scrape_website_data(self, browser, url):
        # The 'scrape_website_data' expects the URL as its second argument after the browser
        return asyncio.get_event_loop().run_until_complete(self.scrape_website_data(browser, url, depth=0, max_depth=3, visited_urls=set()))

    def augment_place_details_with_web_data(self):
        print("Augmenting place details with web data")
        browser = None
        try:
            browser = asyncio.get_event_loop().run_until_complete(launch(headless=False, args=['--no-sandbox', '--disable-setuid-sandbox']))
            for group, places in self.data['grouped_places'].items():
                for place in tqdm(places, desc=f"Processing {group}"):
                    if 'website' in place:
                        web_data = self.sync_scrape_website_data(browser, place['website'])
                        if web_data:  # Check if web_data is not None
                            place.update(web_data)
                        time.sleep(2)
        finally:
            if browser:
                asyncio.get_event_loop().run_until_complete(browser.close())

    # def augment_place_details_with_web_data(self):
    #     print("Augmenting place details with web data")
    #     for group, places in self.data['grouped_places'].items():
    #         for place in tqdm(places, desc=f"Processing {group}"):
    #             if 'website' in place:
    #                 web_data = asyncio.get_event_loop().run_until_complete(self.scrape_website_data(place['website']))
    #                 if web_data:  # Check if web_data is not None
    #                     place.update(web_data)
    #                 time.sleep(2)
                    
    def get_full_link(self, link, base_url):
        if link.startswith('/'):
            return urljoin(base_url, link)
        return link
    
    def is_valid_link(self, link):
        """
        Check if a link is a valid HTTP/HTTPS URL or a valid relative link.
        :param link: Link to check.
        :return: Boolean indicating if the link is valid.
        """
        return link.startswith("http") or (link.startswith("/") and not link.startswith("//"))
    
    ### PRINTING THE FINDINGS TO THE CONSOLE ###
    
    def print_human_readable(self):
        print(f"Research Report for Address: {self.data['address']}\n")
        for group, places in self.data['grouped_places'].items():
            print(f"{group.capitalize()}:")
            for place in places:
                distance = place.get('distance_from_target_address', 'N/A')
                print(f"Distance from Address: {distance} km")
                self.print_place_details(place)
            print("\n")
    
    def print_place_details(self, place):
        # Print basic information
        print(f"Name: {place.get('name', 'N/A')}")
        print(f"Types: {', '.join(place.get('types', ['N/A']))}")
        print(f"Rating: {place.get('rating', 'N/A')}")
        print(f"Price Level: {'$' * place.get('price_level', 0)}")
        
        # Iterate through additional fields
        additional_fields = [
            'formatted_phone_number', 'user_ratings_total', 'website', 'vicinity',
            'formatted_address', 'business_status', 'dine_in', 'editorial_summary',
            'serves_breakfast', 'serves_brunch', 'serves_dinner', 'serves_lunch',
            'serves_wine', 'serves_vegetarian_food', 'takeout'
        ]
        for field in additional_fields:
            value = place.get(field, 'N/A')
            if isinstance(value, list):
                value = ', '.join(value)
            print(f"{field.replace('_', ' ').capitalize()}: {value}")

        # Print links
        print("Menu Links:", ", ".join(place.get('menu_links', [])))
        print("Event Links:", ", ".join(place.get('event_links', [])))
        print("PDF Links:", ", ".join(place.get('pdf_links', [])))
        print("AI Summary:")
        print(place.get('gemini_summary', 'N/A'))
        print("Editorial Summary:")
        print(place.get('editorial_summary', {}).get('overview', 'N/A'))
        print("Hours of Operation:")
        if 'weekday_text' in place.get('opening_hours', {}):
            print('\n'.join(place['opening_hours']['weekday_text']))
        else:
            print("N/A")
        print("Reviews:")
        for review in place.get('reviews', []):
            print(f"{review.get('author_name', 'Anonymous')}: {review.get('text', 'No review text provided')}")
        print("-------------------------------------------------------")
        print()
        print()
    
    ### FILE SAVING METHODS FOR DOCX AND CSV ###

    def save_report_as_word(self, filename):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Open Sans'
        font.size = Pt(12)

        doc.add_heading('Research Report', level=1)
        doc.add_paragraph(f'Address: {self.data["address"]}\n')
        
        self.add_summary_to_doc(doc)

        for group, places in self.data['grouped_places'].items():
            doc.add_heading(f'{group.capitalize()}:', level=2)
            for place in places:
                if isinstance(place, dict):
                    self.add_place_details_to_doc(doc, place)
                else:
                    print(f"Warning: Expected a dictionary for 'place', got {type(place)}")

        doc.save(filename)

    def add_place_details_to_doc(self, doc, place):
        # Ensure 'place' is a dictionary
        if not isinstance(place, dict):
            print(f"Warning: Expected a dictionary for 'place', got {type(place)}")
            return

        doc.add_heading(place.get('name', 'N/A'), level=3)
                
        doc.add_paragraph(f"Distance from Address: {place.get('distance_from_target_address', 'N/A')} km", style='List Bullet')
        
        doc.add_paragraph(f"AI Summary: {place.get('gemini_summary', 'N/A')}", style='List Bullet')
        
        editorial_summary = place.get('editorial_summary', {})
        summary_str = editorial_summary.get('overview', 'N/A')
        doc.add_paragraph(f"Editorial summary: {summary_str}", style='List Bullet')
        
        hours = place.get('opening_hours', {})
        if 'weekday_text' in hours:
            hours_str = '\n'.join(hours['weekday_text'])
        else:
            hours_str = 'N/A'
        doc.add_paragraph(f"Hours of operation: {hours_str}", style='List Bullet')
        
        types = ', '.join(place.get('types', ['N/A']))
        doc.add_paragraph(f"Types: {types}", style='List Bullet')

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Data'
    
        # Standard fields
        standard_fields = ['price_level', 'rating', 'user_ratings_total', 'business_status', 'website', 'url', 'international_phone_number', 'formatted_address']
        for field in standard_fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field.replace('_', ' ').capitalize()
            row_cells[1].text = str(place.get(field, 'N/A'))

        # Additional boolean fields to convert to Yes/No for the report
        additional_fields = ['dine_in', 'takeout', 'reservable', 'serves_breakfast', 'serves_brunch', 'serves_lunch', 'serves_dinner', 'serves_wine', 'serves_vegetarian_food']
        for field in additional_fields:
            value = place.get(field, 'N/A')
            value_text = 'Yes' if value is True else 'No' if value is False else str(value)
            row_cells = table.add_row().cells
            row_cells[0].text = field.replace('_', ' ').capitalize()
            row_cells[1].text = value_text

        # List fields
        self.add_list_to_doc(doc, 'Menu Links', place.get('menu_links', []))
        self.add_list_to_doc(doc, 'Event Links', place.get('event_links', []))
        self.add_list_to_doc(doc, 'PDF Links', place.get('pdf_links', []))
        
        # Gather the reviews into a formatted paragraph and add at the end
        reviews = place.get('reviews', [])
        if reviews:
            doc.add_heading('Reviews', level=4)
            for review in reviews:
                review_text = f"{review.get('author_name', 'Anonymous')}: {review.get('text', 'No review text provided')}"
                doc.add_paragraph(review_text, style='List Bullet')
        else:
            doc.add_paragraph('Reviews: N/A', style='List Bullet')
    
    def add_list_to_doc(self, doc, title, items):
        if items:
            doc.add_paragraph(f'{title}:', style='List Bullet')
            for item in items:
                # Create a sub-bullet point for each item
                p = doc.add_paragraph(style='List Bullet 2')
                p.add_run(item)

    def add_summary_to_doc(self, doc):
        doc.add_heading('Summary of Findings', level=2)
        for group, places in self.data['grouped_places'].items():
            doc.add_heading(f'{group.capitalize()}:', level=3)
            
            # Create a table for this group with an additional column for the weighted score
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Weighted Match Score'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Rating'
            hdr_cells[3].text = 'Price Level'
            hdr_cells[4].text = 'Distance (km)'

            for place in places:
                row_cells = table.add_row().cells
                row_cells[0].text = str(round(place.get('weighted_score', 0), 2))  # Weighted score rounded to 2 decimal places
                row_cells[1].text = place.get('name', 'N/A')
                row_cells[2].text = str(place.get('rating', 'N/A'))
                row_cells[3].text = "$" * place.get('price_level', 0)
                row_cells[4].text = f"{round(place.get('distance_from_target_address', 0), 1)}"

    @staticmethod
    def format_weekday_text(hours):
        return '\n'.join(hours.get('weekday_text', ['N/A'])) if 'weekday_text' in hours else 'N/A'

    @staticmethod
    def format_reviews(reviews):
        formatted_reviews = []
        for review in reviews:
            # Convert Unix timestamp to human-readable date
            review_text = (f"Author: {review.get('author_name', 'Anonymous')}, "
                           f"Rating: {review.get('rating', 'N/A')}, "
                           f"Review: {review.get('text', 'No review text provided')}")
            formatted_reviews.append(review_text)
        return '\n'.join(formatted_reviews) if formatted_reviews else 'N/A'
    
    @staticmethod
    def format_links(links):
        return ', '.join(links) if links else 'N/A'

    @staticmethod
    def save_report_as_csv(json_file_path, csv_file_path):
        with open(json_file_path, 'r') as file:
            data = json.load(file)

        data_for_df = []
        for group, places in data['grouped_places'].items():
            for place in places:
                place_data = {
                    'Group': group.capitalize(),
                    'Name': place.get('name', 'N/A'),
                    'Distance from Address (km)': round(place.get('distance_from_target_address', 0), 1),
                    'Weighted Relevance Score': round(place.get('weighted_score', 0), 2),
                    'Tags': ', '.join(place.get('types', ['N/A'])),
                    'Google Summary': place.get('editorial_summary', {}).get('overview', 'N/A'),
                    'Product Summary': place.get('product_summary', 'N/A'),
                    'AI Summary': place.get('gemini_summary', 'N/A'),
                    'Business Summary': place.get('business_summary', 'N/A'),
                    'Opening Hours': AddressResearcher.format_weekday_text(place.get('opening_hours', {})),
                    'Price Level': place.get('price_level', 'N/A'),
                    'Rating': place.get('rating', 'N/A'),
                    'User Ratings Total': place.get('user_ratings_total', 'N/A'),
                    'Business Status': place.get('business_status', 'N/A'),
                    'Website': place.get('website', 'N/A'),
                    'URL': place.get('url', 'N/A'),
                    'Address': place.get('formatted_address', 'N/A'),
                    'Phone Number': place.get('formatted_phone_number', 'N/A'),
                    'International Phone Number': place.get('international_phone_number', 'N/A'),
                    'Dine In': 'Yes' if place.get('dine_in') else 'No',
                    'Takeout': 'Yes' if place.get('takeout') else 'No',
                    'Reservable': 'Yes' if place.get('reservable') else 'No',
                    'Serves Breakfast': 'Yes' if place.get('serves_breakfast') else 'No',
                    'Serves Lunch': 'Yes' if place.get('serves_lunch') else 'No',
                    'Serves Dinner': 'Yes' if place.get('serves_dinner') else 'No',
                    'Serves Wine': 'Yes' if place.get('serves_wine') else 'No',
                    'Menu Links': AddressResearcher.format_links(place.get('menu_links', [])),
                    'Event Links': AddressResearcher.format_links(place.get('event_links', [])),
                    'PDF Links': AddressResearcher.format_links(place.get('pdf_links', [])),
                    '5 Most Relevant Reviews': AddressResearcher.format_reviews(place.get('reviews', [])),
                    'Menu Content': place.get('menu_content', 'N/A')
                }
                data_for_df.append(place_data)

        df = pd.DataFrame(data_for_df)
        df.to_csv(csv_file_path, index=False)
    
if __name__ == '__main__':
    address_researcher = AddressResearcher(google_maps_api_key)
    address_researcher.perform_research(address)
    address_researcher.augment_place_details_with_web_data()
    address_researcher.print_human_readable()
    address_researcher.save_research_report(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented.json")
    address_researcher.save_report_as_word(f"{FILE_DROP_DIR_PATH}/research_report_doc_formatted.docx")
    # address_researcher.save_report_as_csv(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented.json", f"{FILE_DROP_DIR_PATH}/research_report_spreadsheet.csv")
    AddressResearcher.save_report_as_csv(f"{FILE_DROP_DIR_PATH}/research_report_data_augmented.json", f"{FILE_DROP_DIR_PATH}/research_report_spreadsheet.csv")



